import os
import sys

sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")

import docx

d = docx.Document("Payment interface document-v1—线上支付英文版.docx")

# Walk through body elements in document order
from docx.oxml.ns import qn

body = d.element.body
out_lines = []

# Collect a flat list of (kind, payload) preserving order
elements = []
for child in body.iterchildren():
    tag = child.tag.split("}")[-1]
    if tag == "p":
        # Find corresponding paragraph
        for p in d.paragraphs:
            if p._p is child:
                elements.append(("p", p.text))
                break
    elif tag == "tbl":
        for t in d.tables:
            if t._tbl is child:
                rows = []
                for r in t.rows:
                    rows.append([c.text.strip() for c in r.cells])
                elements.append(("t", rows))
                break

for kind, payload in elements:
    if kind == "p":
        txt = payload.strip()
        if txt:
            out_lines.append(f"P: {txt}")
    else:
        out_lines.append("TABLE:")
        for ri, row in enumerate(payload):
            out_lines.append(f"  R{ri}: " + " | ".join(row))
        out_lines.append("")

with open("docx_ordered.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(out_lines))

print("total elements:", len(elements), flush=True)
print("paragraphs:", sum(1 for k, _ in elements if k == "p"), flush=True)
print("tables:", sum(1 for k, _ in elements if k == "t"), flush=True)
print("output size:", os.path.getsize("docx_ordered.txt"), flush=True)
