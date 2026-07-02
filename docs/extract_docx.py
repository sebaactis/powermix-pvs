import os
import sys

sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")

try:
    files = os.listdir(".")
    print("Files:", files, flush=True)
    fname = next(f for f in files if "Payment" in f)
    print("Found:", repr(fname), flush=True)

    import docx

    d = docx.Document(fname)
    print("paragraphs:", len(d.paragraphs), flush=True)
    print("tables:", len(d.tables), flush=True)

    out = "docx_extracted.txt"
    with open(out, "w", encoding="utf-8") as f:
        f.write("=== PARAGRAPHS ===\n")
        for i, p in enumerate(d.paragraphs):
            txt = p.text.strip()
            if txt:
                f.write(f"[{i}] {txt}\n")
        f.write("\n=== TABLES ===\n")
        for ti, t in enumerate(d.tables):
            f.write(f"--- Table {ti} ---\n")
            for ri, row in enumerate(t.rows):
                cells = [c.text.strip() for c in row.cells]
                f.write(f"  R{ri}: " + " | ".join(cells) + "\n")
    print("saved, size:", os.path.getsize(out), flush=True)
except Exception as e:
    import traceback

    traceback.print_exc()
    sys.exit(1)
